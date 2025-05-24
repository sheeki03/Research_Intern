import streamlit as st
import yaml
import bcrypt
import os
from pathlib import Path
import asyncio
from typing import List, Dict, Any, Optional, Tuple
import io
import re # Added for link extraction
from urllib.parse import urlparse, urljoin # Added for link processing
import pandas as pd # Import pandas

try:
    import fitz  # PyMuPDF
except ImportError:
    # This will be caught at runtime if not installed. 
    # User will need to add PyMuPDF to requirements.txt and rebuild.
    pass 
try:
    from docx import Document # python-docx
except ImportError:
    # This will be caught at runtime if not installed.
    # User will need to add python-docx to requirements.txt and rebuild.
    pass

from src.config import (
    DEBUG,
    LOG_LEVEL,
    OUTPUT_FORMAT,
    OUTPUT_DIR,
    DEFAULT_PROMPTS,
    USERS_CONFIG_PATH,
    SYSTEM_PROMPT,
    OPENROUTER_PRIMARY_MODEL # Import the primary model default
)
# --- DEBUG PRINT 1 --- (immediately after config import)
print(f"DEBUG (main.py top level): config.py OPENROUTER_PRIMARY_MODEL = {OPENROUTER_PRIMARY_MODEL}")

from src.openrouter import OpenRouterClient
from src.firecrawl_client import FirecrawlClient
from src.audit_logger import get_audit_logger, AUDIT_LOG_FILE_PATH # Updated import
from src.core.scanner_utils import discover_sitemap_urls # Added import

# --- Add imports for Chat Models ---
from src.models.chat_models import ChatMessageInput, ChatMessageOutput, ChatSession, ChatHistoryItem
import uuid # For generating session IDs
# --- End imports for Chat Models ---

# --- Add imports for RAG ---
from src.core.rag_utils import (
    get_embedding_model,
    split_text_into_chunks,
    build_faiss_index,
    search_faiss_index,
    DEFAULT_EMBEDDING_MODEL, # Optional: if you want to reference it directly
    TOP_K_RESULTS
)
# --- End imports for RAG ---

# Initialize clients
@st.cache_resource
def init_clients():
    openrouter_client = OpenRouterClient()
    firecrawl_client = FirecrawlClient(
        redis_url=os.getenv("REDIS_URL")
    )
    return openrouter_client, firecrawl_client

def load_users() -> Dict[str, Any]:
    """Load user data from YAML file."""
    if not os.path.exists(USERS_CONFIG_PATH):
        # If no users.yaml, create it with default admin/researcher from init_users.py logic
        # This is a fallback, ideally init_users.py should be run once.
        from src.init_users import init_users as initialize_system_users
        try:
            initialize_system_users()
            st.info("User configuration file not found. Initialized with default users.")
        except Exception as e:
            st.error(f"Failed to initialize default users: {e}")
            return {}
            
    with open(USERS_CONFIG_PATH, 'r') as f:
        users = yaml.safe_load(f)
        if users is None: # Handle empty users.yaml file
            st.warning("User configuration file is empty. Please run user initialization or sign up.")
            return {}
        return users

def save_users(users_data: Dict[str, Any]) -> bool:
    try:
        with open(USERS_CONFIG_PATH, 'w') as f:
            yaml.dump(users_data, f, default_flow_style=False, sort_keys=False)
        return True
    except Exception as e:
        st.error(f"Failed to save user data: {e}")
        return False

def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verify a password against its hash."""
    return bcrypt.checkpw(
        plain_password.encode(),
        hashed_password.encode()
    )

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

# Helper functions for document processing
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF file bytes."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        st.error(f"Error processing PDF: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX file bytes."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error processing DOCX: {e}")
        return ""

def extract_text_from_txt_md(file_bytes: bytes, encoding='utf-8') -> str:
    """Extracts text from TXT or MD file bytes."""
    try:
        return file_bytes.decode(encoding)
    except UnicodeDecodeError:
        try:
            return file_bytes.decode('latin-1') # Fallback encoding
        except Exception as e:
            st.error(f"Error decoding text file with fallback: {e}")
            return ""
    except Exception as e:
        st.error(f"Error processing text/markdown file: {e}")
        return ""
# End of helper functions

async def process_urls(urls: List[str], client: FirecrawlClient) -> List[Dict[str, Any]]:
    """Process a list of specific URLs and return a list of dicts with url, content/error."""
    results = await client.scrape_multiple_urls(urls)
    
    processed_results = []
    for result in results:
        url = result.get("metadata", {}).get("url", result.get("url", "unknown URL"))
        if result.get("success", False):
            content_data = result.get("data", {}).get("content", "")
            if not content_data:
                 content_data = result.get("content", "")
            processed_results.append({"url": url, "content": content_data, "status": "success"})
        else:
            error_message = result.get("error", "Unknown error")
            processed_results.append({"url": url, "error": error_message, "status": "failed"})
            
    return processed_results

async def perform_web_research(query: str, client: OpenRouterClient) -> str:
    """Perform web research on a given query."""
    # Generate search queries
    queries = await client.generate_serp_queries(query)
    
    # Simulate search results (replace with actual search implementation)
    search_results = {
        "data": [
            {"content": "Sample content 1", "url": "https://example.com/1"},
            {"content": "Sample content 2", "url": "https://example.com/2"}
        ]
    }
    
    # Process search results
    results = await client.process_serp_result(query, search_results)
    
    # Write final report
    report = await client.write_final_report(
        query,
        results["learnings"],
        [item["url"] for item in search_results["data"]]
    )
    
    return report

# --- Add Model Definitions Here ---
# Define model choices based on user query and web search results
_MODEL_OPTIONS_RAW = {
    # "Mistral Medium 3": "mistralai/mistral-medium-3", # Replaced
    "Qwen3 30B": "qwen/qwen3-30b-a3b:free", # Added new model
    "Google Gemini 2.5 Pro": "google/gemini-2.5-pro-preview",
    "OpenAI o3": "openai/o3",
    "OpenAI GPT-4.1": "openai/gpt-4.1",
    "Claude 3.7 Sonnet (thinking)": "anthropic/claude-3.7-sonnet:thinking",
    "DeepSeek R1T Chimera": "tngtech/deepseek-r1t-chimera:free",
}

MODEL_OPTIONS = {}
TEMP_MODEL_DISPLAY_NAMES = []
for original_key, identifier in _MODEL_OPTIONS_RAW.items():
    cleaned_key = re.sub(r'\s*\([Ff]ree\)$', '', original_key).strip()
    MODEL_OPTIONS[cleaned_key] = identifier
    TEMP_MODEL_DISPLAY_NAMES.append(cleaned_key)

MODEL_DISPLAY_NAMES = TEMP_MODEL_DISPLAY_NAMES
MODEL_IDENTIFIERS = list(MODEL_OPTIONS.values())
# --- End Model Definitions ---

# --- NEW: Crawl and Scrape Function --- 
async def crawl_and_scrape_site(start_url: str, limit: int, client: FirecrawlClient) -> List[Dict[str, Any]]:
    """Crawls a website starting from start_url, scraping content and following same-domain links up to a limit."""
    if not start_url or not client.validate_url(start_url):
        st.warning(f"Invalid start URL provided for crawl: {start_url}")
        return []

    if limit <= 0:
        st.warning("Crawl limit must be greater than 0.")
        return []

    base_domain = urlparse(start_url).netloc
    urls_to_scrape = {start_url}
    visited_urls = set()
    scraped_data_list = []
    scrape_count = 0

    # Modified regex to use r"..." string format
    link_regex = re.compile(r'href\s*=\s*[\'"]([^\'"]+)[\'"]', re.IGNORECASE)

    st.info(f"Starting crawl from {start_url} (Domain: {base_domain}), limit: {limit} pages.")

    while urls_to_scrape and scrape_count < limit:
        current_url = urls_to_scrape.pop()

        if current_url in visited_urls:
            continue

        visited_urls.add(current_url)
        scrape_count += 1
        st.write(f"Crawling [{scrape_count}/{limit}]: {current_url}") # Show progress

        try:
            # Use the single scrape_url method from the client
            scraped_result_dict = await client.scrape_url(current_url)
            
            # Extract markdown content for storage/AI processing
            markdown_content = scraped_result_dict.get("data", {}).get("content", "") 
            # Extract HTML content specifically for link extraction
            html_for_links = scraped_result_dict.get("data", {}).get("html_content", "")
            
            # Check for errors reported by the client
            client_error = scraped_result_dict.get("error")
            if client_error:
                st.error(f"Scraper client returned error for {current_url}: {client_error}")
                scraped_data_list.append({"url": current_url, "error": client_error, "status": "failed"})
                get_audit_logger( # Replaced log_audit_event
                    user=st.session_state.get('username', 'SYSTEM'),
                    role=st.session_state.get('role', 'N/A'),
                    action="CRAWL_SCRAPE_CLIENT_ERROR", 
                    details=f"Firecrawl client processing error for {current_url}: {client_error}",
                    links=[current_url]
                )
                continue # Move to the next URL

            if markdown_content:
                # Store the markdown content
                scraped_data_list.append({"url": current_url, "content": markdown_content, "status": "success"})
                
                # Find and process links using HTML content if available, otherwise fallback to markdown
                link_source_content = html_for_links if html_for_links else markdown_content
                if scrape_count < limit and link_source_content:
                    found_links_count = 0
                    potential_links = link_regex.findall(link_source_content)
                    for link in potential_links:
                        try:
                            absolute_link = urljoin(current_url, link.strip())
                            parsed_link = urlparse(absolute_link)
                            
                            # Basic validation and domain check
                            if parsed_link.scheme in ['http', 'https'] and parsed_link.netloc == base_domain:
                                if absolute_link not in visited_urls and absolute_link not in urls_to_scrape:
                                    urls_to_scrape.add(absolute_link)
                                    found_links_count += 1
                        except Exception as link_e:
                            # Ignore errors parsing/resolving individual links
                            # print(f" Minor error processing link '{link}': {link_e}")
                            pass # Be less verbose in UI
                    # st.write(f"  Found {found_links_count} new links on {current_url}") # Optional debug
            else: # No markdown content extracted
                error_msg = f"No primary (markdown) content extracted from: {current_url}"
                scraped_data_list.append({"url": current_url, "error": error_msg, "status": "no_content"})
                st.warning(error_msg)
        
        except Exception as e: # Catch exceptions from scrape_url call itself or processing
            error_msg = f"Error scraping {current_url}: {str(e)}"
            st.error(error_msg)
            scraped_data_list.append({"url": current_url, "error": error_msg, "status": "failed"})
            # Log this significant error
            get_audit_logger( # Replaced log_audit_event
                user=st.session_state.get('username', 'SYSTEM'),
                role=st.session_state.get('role', 'N/A'),
                action="CRAWL_SCRAPE_FAILURE",
                details=error_msg,
                links=[current_url]
            )

    if urls_to_scrape:
        st.warning(f"Crawl limit ({limit}) reached, {len(urls_to_scrape)} URLs remaining in queue.")
    
    successful_scrape_count = sum(1 for item in scraped_data_list if item.get('status') == 'success' and item.get('content'))
    st.info(f"Crawl finished. Attempted {len(scraped_data_list)} pages, successfully scraped {successful_scrape_count} with content.")
    return scraped_data_list
# --- End Crawl Function --- 

# --- Add Function to Parse Log Line ---
def parse_log_line(line: str) -> Optional[Dict[str, str]]:
    """Parses a single line from the audit log file."""
    parts = line.strip().split(' | ')
    if len(parts) < 4: # Basic check for minimum parts (Timestamp, User, Role, Action)
        return None
    
    log_entry = {
        "Timestamp": parts[0],
        "Username": "N/A",
        "Role": "N/A",
        "Action": "N/A",
        "Model": "N/A",
        "Links": "N/A",
        "Details": "N/A"
    }
    
    for part in parts[1:]: # Skip timestamp
        try:
            key, value = part.split(': ', 1)
            # Map keys, handling potential variations
            if key == "USER": log_entry["Username"] = value
            elif key == "ROLE": log_entry["Role"] = value
            elif key == "ACTION": log_entry["Action"] = value
            elif key == "MODEL": log_entry["Model"] = value
            elif key == "LINKS": log_entry["Links"] = value
            elif key == "DETAILS": log_entry["Details"] = value
        except ValueError:
             # If a part doesn't split correctly, add it to Details as unparsed info
             log_entry["Details"] = f"{log_entry.get('Details', '')} [Unparsed: {part}]"
            
    return log_entry
# --- End Function ---

# --- Chat UI / Logic Functions ---

def get_or_create_chat_session(report_id: str, session_id: Optional[str] = None) -> ChatSession:
    """Retrieves an existing chat session from st.session_state or creates a new one."""
    if "chat_sessions_store" not in st.session_state:
        st.session_state.chat_sessions_store = {}

    if session_id and session_id in st.session_state.chat_sessions_store:
        session = st.session_state.chat_sessions_store[session_id]
        if session.report_id == report_id:
            return session
        # else: if session_id is for a different report, we should create a new one for this report_id
        # This logic implies that session_id might be tied to a report view, not globally unique across reports.
        # For simplicity, if a session_id is passed but for wrong report, we create a new one for current report.

    # Create a new session for the current report_id
    new_session = ChatSession(report_id=report_id) # ChatSession model generates its own session_id
    st.session_state.chat_sessions_store[new_session.session_id] = new_session
    # Store the new session_id as the active one for the current chat interface context
    st.session_state.current_chat_session_id = new_session.session_id 
    return new_session

async def handle_chat_submission(user_query: str, report_id: str, current_session_id: Optional[str]) -> Tuple[str, str]:
    """Handles user chat submission, gets LLM response using RAG, updates history, returns AI response and session_id."""
    session = get_or_create_chat_session(report_id=report_id, session_id=current_session_id)
    session.history.append(ChatHistoryItem(role="user", content=user_query))

    # --- AI Logic (Task 10: RAG Integration) ---
    ai_response_content = "Sorry, I encountered an error trying to process your question with RAG."
    retrieved_context_for_llm = "No specific context retrieved from RAG for this query."

    try:
        embedding_model = get_embedding_model() # Cached model
        rag_data_for_report = st.session_state.get("rag_contexts", {}).get(report_id)

        if rag_data_for_report and rag_data_for_report.get("index") is not None and rag_data_for_report.get("chunks") is not None:
            faiss_index = rag_data_for_report["index"]
            text_chunks = rag_data_for_report["chunks"]
            
            # st.write(f"DEBUG: RAG index found for report {report_id}. Chunks: {len(text_chunks)}. Index items: {faiss_index.ntotal if faiss_index else 'N/A'}")
            
            with st.spinner("Searching relevant context for your query..."):
                relevant_chunks = search_faiss_index(
                    index=faiss_index, 
                    query_text=user_query, 
                    embedding_model=embedding_model, 
                    text_chunks=text_chunks, 
                    top_k=TOP_K_RESULTS
                )
            
            if relevant_chunks:
                retrieved_context_for_llm = "\n\n---\n\n".join(relevant_chunks)
                # Optional: Log or display what was retrieved for debugging
                # st.sidebar.expander(f"Retrieved RAG Chunks for '{user_query[:30]}...'").write(relevant_chunks)
            else:
                retrieved_context_for_llm = "No highly relevant information found in the knowledge base for your specific query. I will answer based on general knowledge and conversation history."
        elif rag_data_for_report is None: # Explicitly None means RAG building failed or no content
             retrieved_context_for_llm = "RAG context was not successfully built for this report (e.g. no content or error during indexing). Answering based on general knowledge and history."
             st.warning(f"RAG context not available or failed to build for report {report_id}. Chat accuracy might be reduced.")
        else: # Should not happen if rag_contexts[report_id] is always set to dict or None
            retrieved_context_for_llm = "RAG context store seems to be in an unexpected state. Answering based on general knowledge and history."
            st.error(f"Unexpected RAG context state for report {report_id}.")

        # Format conversation history
        formatted_history = ""
        MAX_HISTORY_TURNS = 5 
        relevant_history = session.history[:-1] 
        start_index = max(0, len(relevant_history) - (2 * MAX_HISTORY_TURNS))
        for msg in relevant_history[start_index:]:
            role_display = "User" if msg.role == "user" else "AI"
            formatted_history += f"{role_display}: {msg.content}\n"

        llm_prompt = (
            f"You are a helpful AI assistant. You have been provided with conversation history and specific, highly relevant excerpts from a larger knowledge base (retrieved via RAG). "
            f"Please answer the user's current question based *primarily* on these retrieved excerpts and the conversation history. "
            f"If the retrieved excerpts are insufficient or state that no relevant information was found, you may say so or use general knowledge cautiously.\n\n"
        )

        if formatted_history:
            llm_prompt += f"CONVERSATION HISTORY:\n---\n{formatted_history}---\n\n"
        
        llm_prompt += f"RETRIEVED CONTEXTUAL EXCERPTS (RAG):\n---\n{retrieved_context_for_llm}\n---\n\n"
            
        llm_prompt += (
            f"USER'S CURRENT QUESTION: {user_query}\n\n"
            f"Based on the retrieved excerpts and conversation history, YOUR ANSWER:"
        )
        
        if "openrouter_client" not in st.session_state:
             st.session_state.openrouter_client, _ = init_clients()

        if st.session_state.openrouter_client:
            chat_system_prompt = (
                "You are an AI assistant. Answer the user's question based on the provided conversation history and the highly relevant contextual excerpts retrieved from a knowledge base."
            )
            
            # --- DEBUG PRINT ADDED ---
            model_to_use_for_llm = st.session_state.get("selected_model", OPENROUTER_PRIMARY_MODEL)
            print(f"DEBUG: About to call LLM. Model identifier type: {type(model_to_use_for_llm)}, value: {model_to_use_for_llm}")
            # --- END DEBUG PRINT ---

            # --- CHATBOX MODEL OVERRIDE ---
            # For chat, always use the specified chimera model, overriding user selection for report gen.
            chat_specific_model = "tngtech/deepseek-r1t-chimera:free" # Can also use OPENROUTER_PRIMARY_MODEL if it's guaranteed to be this.
            print(f"DEBUG: Chatbox overriding LLM to: {chat_specific_model}")
            # --- END CHATBOX MODEL OVERRIDE ---

            generated_text = await st.session_state.openrouter_client.generate_response(
                prompt=llm_prompt,
                system_prompt=chat_system_prompt, 
                model_override=chat_specific_model # Use the chat-specific model
            )
            if generated_text:
                ai_response_content = generated_text
            else:
                ai_response_content = "I received an empty response from the AI. Please try rephrasing your question."
        else:
            ai_response_content = "AI client not available. Cannot process question."

    except Exception as e:
        st.error(f"Error during RAG processing or LLM call: {str(e)}")
        ai_response_content = f"Sorry, I encountered an error: {str(e)}"
    # --- End AI Logic ---
    
    session.history.append(ChatHistoryItem(role="ai", content=ai_response_content))
    st.session_state.chat_sessions_store[session.session_id] = session
    return ai_response_content, session.session_id

async def display_chat_interface_v2():
    # st.subheader("Chat with AI about this Report") # Title will be in expander

    report_id = st.session_state.get("current_report_id_for_chat")
    if not report_id:
        # This case should ideally not be reached if expander visibility is also tied to report_id presence
        st.warning("Chat cannot be initialized: No active report ID found.")
        return

    with st.expander("💬 Chat with AI about this Report", expanded=st.session_state.get("chat_ui_expanded", False)):
        current_chat_session_id = st.session_state.get("current_chat_session_id")
        active_session = get_or_create_chat_session(report_id, current_chat_session_id)
        st.session_state.current_chat_session_id = active_session.session_id

        # Message display area within the expander
        # To make it scrollable and limit height, we can wrap messages in a container with a fixed height
        # However, st.container() itself doesn't take height. We might need custom CSS or just let it expand.
        # For now, let it expand naturally within the expander.
        for msg in active_session.history:
            with st.chat_message(msg.role):
                st.markdown(msg.content)

        if st.session_state.get("ai_is_thinking", False):
            with st.chat_message("assistant"):
                st.markdown("_AI is thinking..._")

        user_prompt = st.chat_input(f"Ask about report {report_id}...", key=f"chat_input_{report_id}") # Unique key for chat input

        if user_prompt and not st.session_state.get("ai_is_thinking", False):
            st.session_state.last_user_prompt_for_processing = user_prompt
            st.session_state.ai_is_thinking = True
            st.rerun()
        elif st.session_state.get("ai_is_thinking", False) and "last_user_prompt_for_processing" in st.session_state:
            prompt_to_process = st.session_state.pop("last_user_prompt_for_processing", None)
            if prompt_to_process:
                await handle_chat_submission( 
                    user_query=prompt_to_process,
                    report_id=report_id,
                    current_session_id=active_session.session_id
                )
            st.session_state.ai_is_thinking = False
            st.rerun()

# --- END Chat UI / Logic Functions ---

async def main():
    st.set_page_config(
        page_title="Research Agent",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    # --- DEBUG PRINT 2 --- (start of main() async function)
    print(f"DEBUG (main() start): config.py OPENROUTER_PRIMARY_MODEL = {OPENROUTER_PRIMARY_MODEL}")

    # Initialize chat messages in session state if not already there
    # This is a good place for it, or right after page_config
    if "chat_messages" not in st.session_state:
        st.session_state.chat_messages = []
    if "report_generated_for_chat" not in st.session_state: # Flag to control chat display
        st.session_state.report_generated_for_chat = False
    if "current_report_id_for_chat" not in st.session_state: # To store which report is active
        st.session_state.current_report_id_for_chat = None
    if "chat_sessions_store" not in st.session_state: # To store all chat sessions
        st.session_state.chat_sessions_store = {}
    if "current_chat_session_id" not in st.session_state: # Active session ID for current report
        st.session_state.current_chat_session_id = None
    if "openrouter_client" not in st.session_state: # Store clients in session_state
        st.session_state.openrouter_client, st.session_state.firecrawl_client = init_clients()
    if "ai_is_thinking" not in st.session_state: # Changed here for initialization
        st.session_state.ai_is_thinking = False
    if "last_user_prompt_for_processing" not in st.session_state: 
        st.session_state.last_user_prompt_for_processing = None
    if "rag_contexts" not in st.session_state: # For storing RAG indexes and chunks per report
        st.session_state.rag_contexts = {}
    if "chat_ui_expanded" not in st.session_state: # For controlling chat expander
        st.session_state.chat_ui_expanded = False

    # Initialize session state for system_prompt if not already set
    if "system_prompt" not in st.session_state:
        st.session_state.system_prompt = SYSTEM_PROMPT # Global default
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "username" not in st.session_state:
        st.session_state.username = None
    if "show_signup" not in st.session_state:
        st.session_state.show_signup = False
    if "processed_documents_content" not in st.session_state: 
        st.session_state.processed_documents_content = []
    if "last_uploaded_file_details" not in st.session_state: # For comparing if files changed
        st.session_state.last_uploaded_file_details = []
    if "unified_report_content" not in st.session_state: 
        st.session_state.unified_report_content = ""
    if "scraped_web_content" not in st.session_state: # Added for this task (Task 3)
        st.session_state.scraped_web_content = []
    if "crawled_web_content" not in st.session_state:
        st.session_state.crawled_web_content = []
    # Session state for sitemap scanning (Task 2.2)
    if "discovered_sitemap_urls" not in st.session_state:
        st.session_state.discovered_sitemap_urls = []
    if "sitemap_scan_in_progress" not in st.session_state:
        st.session_state.sitemap_scan_in_progress = False
    if "sitemap_scan_error" not in st.session_state:
        st.session_state.sitemap_scan_error = None
    if "sitemap_scan_completed" not in st.session_state: # To know if a scan has been attempted
        st.session_state.sitemap_scan_completed = False
    if "selected_sitemap_urls" not in st.session_state: # For Task 2.3
        st.session_state.selected_sitemap_urls = set() # Using a set for selected URLs

    # Sidebar
    with st.sidebar:
        st.title("Access Panel")
        
        if not st.session_state.authenticated:
            if st.session_state.show_signup:
                st.subheader("Create Account")
                new_username = st.text_input("New Username", key="signup_username")
                new_password = st.text_input("New Password", type="password", key="signup_password")
                confirm_password = st.text_input("Confirm Password", type="password", key="signup_confirm_password")
                
                if st.button("Sign Up", key="signup_submit_button"):
                    if not new_username or not new_password:
                        st.error("Username and password cannot be empty.")
                        get_audit_logger( # Replaced log_audit_event
                            user=new_username or "ANONYMOUS", 
                            role="N/A", 
                            action="USER_SIGNUP_FAILURE", 
                            details="Attempted signup with empty username/password."
                        )
                    elif new_password != confirm_password:
                        st.error("Passwords do not match.")
                        get_audit_logger( # Replaced log_audit_event
                            user=new_username, 
                            role="N/A", 
                            action="USER_SIGNUP_FAILURE", 
                            details="Attempted signup with non-matching passwords."
                        )
                    else:
                        users = load_users()
                        if new_username in users:
                            st.error("Username already exists.")
                            get_audit_logger( # Replaced log_audit_event
                                user=new_username, 
                                role="N/A", 
                                action="USER_SIGNUP_FAILURE", 
                                details=f"Attempted signup with existing username: {new_username}."
                            )
                        else:
                            users[new_username] = {
                                "password": hash_password(new_password),
                                "role": "researcher", # Default role
                                "system_prompt": DEFAULT_PROMPTS.get("researcher", SYSTEM_PROMPT)
                            }
                            if save_users(users):
                                st.success("Account created successfully! Please log in.")
                                get_audit_logger( # Replaced log_audit_event
                                    user=new_username, 
                                    role="researcher", 
                                    action="USER_SIGNUP_SUCCESS", 
                                    details=f"New user account created: {new_username}"
                                )
                                st.session_state.show_signup = False # Switch back to login
                            else:
                                st.error("Failed to create account. Please try again.")
                                get_audit_logger( # Replaced log_audit_event
                                    user=new_username, 
                                    role="researcher", 
                                    action="USER_SIGNUP_FAILURE", 
                                    details="Failed to save new user account: {new_username} after validation."
                                )
                
                if st.button("Back to Login", key="back_to_login_button"):
                    st.session_state.show_signup = False
                    st.rerun()
            else:
                st.subheader("Login")
                username_input = st.text_input("Username", key="login_username")
                password_input = st.text_input("Password", type="password", key="login_password")
                
                if st.button("Login", key="login_button"):
                    users = load_users()
                    user_data = users.get(username_input, {})
                    if user_data and verify_password(password_input, user_data.get("password", "")):
                        st.session_state.authenticated = True
                        st.session_state.username = username_input
                        # Load user-specific prompt, fallback to role-based, then global default
                        user_specific_prompt = user_data.get("system_prompt")
                        current_role = user_data.get("role", "researcher") # Default role if not specified
                        st.session_state.role = current_role # Store role in session state

                        if not user_specific_prompt:
                            user_specific_prompt = DEFAULT_PROMPTS.get(current_role, SYSTEM_PROMPT)
                        
                        st.session_state.system_prompt = user_specific_prompt
                        st.success("Login successful!")
                        get_audit_logger( # Replaced log_audit_event
                            user=username_input, 
                            role=current_role, 
                            action="USER_LOGIN_SUCCESS", 
                            details=f"User {username_input} logged in successfully."
                        )
                        st.rerun()
                    else:
                        st.error("Invalid username or password")
                        get_audit_logger( # Replaced log_audit_event
                            user=username_input or "UNKNOWN_USER", 
                            role="N/A", 
                            action="USER_LOGIN_FAILURE", 
                            details=f"Failed login attempt for username: '{username_input}'."
                        )
                
                if st.button("Create Account", key="show_signup_button"):
                    st.session_state.show_signup = True
                    st.rerun()
        else:
            st.write(f"Logged in as: {st.session_state.username}")
            if st.button("Logout", key="logout_button"):
                logged_out_username = st.session_state.username
                logged_out_role = st.session_state.get("role", "N/A")
                st.session_state.authenticated = False
                st.session_state.username = None
                st.session_state.role = None # Clear role on logout
                st.session_state.system_prompt = SYSTEM_PROMPT # Reset to global default on logout
                st.session_state.show_signup = False # Reset signup view on logout
                get_audit_logger( # Replaced log_audit_event
                    user=logged_out_username, 
                    role=logged_out_role, 
                    action="USER_LOGOUT", 
                    details=f"User {logged_out_username} logged out."
                )
                st.rerun()

            st.markdown("---") # Separator
            st.subheader("Session System Prompt")
            # Allow editing of the session system prompt
            new_prompt = st.text_area(
                "Edit current session's system prompt:", 
                value=st.session_state.system_prompt, 
                height=300,
                key="session_prompt_editor"
            )
            if new_prompt != st.session_state.system_prompt:
                st.session_state.system_prompt = new_prompt
                st.success("Session system prompt updated.")
                get_audit_logger( # Replaced log_audit_event
                    user=st.session_state.username, 
                    role=st.session_state.get("role", "N/A"), 
                    action="SESSION_PROMPT_UPDATED", 
                    details=f"User updated session prompt. New prompt: '{new_prompt}'"
                )

    # Main content area
    if st.session_state.authenticated:
        st.title("AI Research Agent")
        # Clients are now in session_state, no need to unpack here unless used directly in main's top level scope
        # openrouter_client, firecrawl_client = st.session_state.openrouter_client, st.session_state.firecrawl_client

        st.header("Unified Research Interface")

        # --- Add Model Selection Section ---
        st.subheader("Model Selection")
        # Determine the index of the currently configured primary model for default selection
        try:
            default_model_identifier = OPENROUTER_PRIMARY_MODEL # Loaded from config/env
            # --- DEBUG PRINT 3 --- (before determining selectbox default)
            print(f"DEBUG (selectbox default logic): default_model_identifier for selectbox = {default_model_identifier}")
            default_display_name_for_default_identifier = ""
            # Find the cleaned display name that maps to the default identifier
            for display_name, identifier_in_map in MODEL_OPTIONS.items(): 
                if identifier_in_map == default_model_identifier:
                    default_display_name_for_default_identifier = display_name
                    break
            
            if default_display_name_for_default_identifier:
                default_index = MODEL_DISPLAY_NAMES.index(default_display_name_for_default_identifier)
            else: # Fallback if the identifier is not found
                st.warning(f"Default primary model identifier '{default_model_identifier}' not found in the available MODEL_OPTIONS mapping. Defaulting to first option.")
                default_index = 0 
        except ValueError: # Handles if .index() fails for some reason
            default_index = 0 
            st.warning(f"Could not determine default index for model '{OPENROUTER_PRIMARY_MODEL}'. Defaulting to first option.")

        selected_model_display_name = st.selectbox(
            "Choose the AI model for report generation:",
            options=MODEL_DISPLAY_NAMES, # Use the cleaned list
            index=default_index,
            key="model_selector",
            help="Select the AI model to use. Defaults are set in config/env."
        )
        # Store the corresponding identifier in session state using the cleaned display name for lookup
        st.session_state.selected_model = MODEL_OPTIONS[selected_model_display_name]
        st.markdown("---") # Separator
        # --- End Model Selection Section ---

        # Research Query Input
        st.subheader("1. Define Your Research Focus (Optional)")
        research_query = st.text_area(
            "Enter your research query or specific questions:",
            height=100,
            key="research_query_input",
            help="Clearly state what you want the AI to investigate or analyze based on the provided documents and URLs."
        )

        # Document Upload Section
        st.subheader("2. Upload Relevant Documents (Optional)")
        
        uploaded_files_new = st.file_uploader(
            "Upload documents (PDF, DOCX, TXT, MD)",
            type=["pdf", "docx", "txt", "md"],
            accept_multiple_files=True,
            key="unified_file_uploader", # Using a consistent key might be okay if Streamlit handles updates well
            help="Upload whitepapers, reports, articles, or any other relevant documents."
        )

        # Check if the set of uploaded files has actually changed
        current_file_details = [(f.name, f.size) for f in uploaded_files_new] if uploaded_files_new else []
        files_have_changed = (current_file_details != st.session_state.get("last_uploaded_file_details", []))

        if uploaded_files_new and files_have_changed:
            st.session_state.last_uploaded_file_details = current_file_details
            st.session_state.processed_documents_content = [] # Clear previous batch to process anew
            current_batch_processed_content = []
            
            # This inner check is a bit redundant given the outer `if uploaded_files_new` but safe.
            # if uploaded_files_new: 
            with st.status(f"Processing {len(uploaded_files_new)} uploaded file(s)...", expanded=True) as status_container:
                # progress_bar = st.progress(0) # Using st.status handles ongoing updates better for this pattern
                
                for i, uploaded_file_data in enumerate(uploaded_files_new):
                    st.write(f"Processing: {uploaded_file_data.name} ({i+1}/{len(uploaded_files_new)})")
                    # action_details = f"File upload attempt: {uploaded_file_data.name}" # Audit log moved
                    # log_status = "SUCCESS"

                    file_bytes = uploaded_file_data.getvalue()
                    content = ""
                    # Determine file type for processing (prefer extension for clarity)
                    file_extension = uploaded_file_data.name.split('.')[-1].lower()

                    if file_extension == "pdf":
                        content = extract_text_from_pdf(file_bytes)
                    elif file_extension == "docx":
                        content = extract_text_from_docx(file_bytes)
                    elif file_extension in ["txt", "md"]:
                        content = extract_text_from_txt_md(file_bytes)
                    else:
                        st.warning(f"Unsupported file type skipped (extension not matched): {uploaded_file_data.name}")
                        content = "" 
                        # log_status = "SKIPPED_UNSUPPORTED"

                    if content:
                        current_batch_processed_content.append({"name": uploaded_file_data.name, "text": content})
                        st.success(f"Successfully extracted text from: {uploaded_file_data.name}")
                        action_details = f"Successfully extracted content from uploaded file: {uploaded_file_data.name}"
                        get_audit_logger( # Replaced log_audit_event
                            user=st.session_state.username, 
                            role=st.session_state.get('role','N/A'), 
                            action="FILE_PROCESS_SUCCESS", 
                            details=action_details
                        )
                    elif file_extension in ["pdf", "docx", "txt", "md"]:
                        st.error(f"Failed to extract text from: {uploaded_file_data.name} (see specific error above if any).")
                        action_details = f"Failed to extract content from uploaded file: {uploaded_file_data.name}"
                        get_audit_logger( # Replaced log_audit_event
                            user=st.session_state.username, 
                            role=st.session_state.get('role','N/A'), 
                            action="FILE_PROCESS_FAILURE", 
                            details=action_details
                        )
                    else: # If skipped due to unsupported type and no content
                        action_details = f"Skipped unsupported file type: {uploaded_file_data.name}"
                        get_audit_logger( # Replaced log_audit_event
                            user=st.session_state.username, 
                            role=st.session_state.get('role','N/A'), 
                            action="FILE_PROCESS_SKIPPED", 
                            details=action_details
                        )

                    # progress_bar.progress((i + 1) / len(uploaded_files_new)) # Handled by st.status context
                
                st.session_state.processed_documents_content = current_batch_processed_content
                status_container.update(label=f"{len(current_batch_processed_content)} out of {len(uploaded_files_new)} files processed successfully.", state="complete", expanded=False)
                # Rerun to update displays based on newly processed files in session_state
                # This helps ensure UI consistency immediately after processing
                st.rerun() 
        elif not uploaded_files_new: # If no files are selected (uploader is cleared)
            if st.session_state.last_uploaded_file_details: # If there were files before, clear details
                st.session_state.last_uploaded_file_details = []
                st.session_state.processed_documents_content = []
                st.rerun() # Rerun to reflect that no files are processed

        # Display summary of currently processed documents from session_state
        if st.session_state.get("processed_documents_content"):
            st.markdown("---")
            st.subheader(f"Processed Documents ({len(st.session_state.processed_documents_content)} ready for report):")
            for doc_info in st.session_state.processed_documents_content:
                with st.expander(f"{doc_info['name']} ({len(doc_info['text'])} chars) - Click to preview"):
                    st.caption(doc_info['text'][:250] + "..." if len(doc_info['text']) > 250 else doc_info['text'])
            st.markdown("---")

        # URL Input Section
        st.subheader("3. Provide Specific Web URLs (Optional)")
        # st.write("Enter up to 10 URLs for web content scraping:") # Old instruction
        
        # Create a list to hold URL inputs (Old)
        # url_inputs = [] 
        # for i in range(10):
        #     url = st.text_input(f"URL {i+1}", key=f"url_input_{i+1}", placeholder=f"https://example.com/page{i+1}")
        #     url_inputs.append(url)

        # New Text Area for URLs
        urls_text_area = st.text_area(
            "Enter URLs, one per line:",
            height=150, # Adjust height as needed
            key="urls_text_area_input",
            placeholder="https://example.com/page1\nhttps://another-example.com/article2\n..."
        )
        
        # Collect provided URLs from the text area
        # submitted_urls = [url for url in url_inputs if url] # Old collection logic
        if urls_text_area:
            submitted_urls = [url.strip() for url in urls_text_area.split('\n') if url.strip()] 
        else:
            submitted_urls = []

        # --- Add Crawl Input Section --- (Header 4)
        st.subheader("4. Crawl & Scrape Site (Optional)")
        st.markdown("""
This section offers two methods to gather content from a website. Choose one option if you wish to use these features:

- **Option 4a (Scan & Select):** Scans a site's sitemap to list its URLs, allowing you to select specific pages for scraping. This is useful when you want to target particular pages from a known list.
- **Option 4b (Crawl & Discover):** Starts at a given URL and automatically crawls the site by following links to other pages on the same website, scraping their content up to a specified limit. This is useful for broader content discovery when you don't have a specific list of pages.

If you use **Option 4a**, the URLs you select will be used for scraping when you generate the report. 
If you use **Option 4b**, the crawled pages will be used. 
If specific URLs are provided via sitemap selection (4a) or manual entry in section 3, crawling (4b) will be skipped even if a start URL is provided for it.
""")
        
        st.markdown("**Option 4a: Scan Site for Specific URLs from Sitemap**")
        site_to_scan_url = st.text_input(
            "URL to Scan for Sitemaps:",
            key="site_to_scan_url_input",
            placeholder="https://example.com (domain or full URL)",
            help="Enter a website URL to scan its sitemap for a list of all its pages. You can then select which pages to scrape."
        )
        # Placeholder for Scan Site button logic (Task 2.2)
        if st.button("Scan Site for URLs", key="scan_site_button", disabled=st.session_state.sitemap_scan_in_progress):
            if site_to_scan_url:
                st.session_state.sitemap_scan_in_progress = True
                st.session_state.discovered_sitemap_urls = [] # Clear previous results
                st.session_state.sitemap_scan_error = None
                st.session_state.sitemap_scan_completed = False # Reset completed flag
                
                # Log scan initiation
                get_audit_logger(
                    user=st.session_state.username,
                    role=st.session_state.get('role', 'N/A'),
                    action="SITEMAP_SCAN_INITIATED",
                    details=f"Sitemap scan initiated for URL: {site_to_scan_url}"
                )

                try:
                    # Use a spinner for better UX during async operation
                    with st.spinner(f"Scanning {site_to_scan_url} for sitemap URLs... Please wait."):
                        discovered_urls = await discover_sitemap_urls(site_to_scan_url)
                    
                    st.session_state.discovered_sitemap_urls = discovered_urls
                    st.session_state.sitemap_scan_completed = True
                    if discovered_urls:
                        st.success(f"Scan complete! Found {len(discovered_urls)} URLs.")
                        get_audit_logger(
                            user=st.session_state.username,
                            role=st.session_state.get('role', 'N/A'),
                            action="SITEMAP_SCAN_SUCCESS",
                            details=f"Sitemap scan successful for {site_to_scan_url}. Found {len(discovered_urls)} URLs."
                        )
                    else:
                        st.info(f"Scan complete. No URLs found for {site_to_scan_url} via sitemaps.")
                        get_audit_logger(
                            user=st.session_state.username,
                            role=st.session_state.get('role', 'N/A'),
                            action="SITEMAP_SCAN_NO_URLS_FOUND",
                            details=f"Sitemap scan completed for {site_to_scan_url}, but no URLs were found."
                        )
                except Exception as e:
                    st.session_state.sitemap_scan_error = f"An error occurred during sitemap scan: {str(e)}"
                    st.error(st.session_state.sitemap_scan_error)
                    st.session_state.sitemap_scan_completed = True # Mark as completed even on error to stop spinner logic if page reloads
                    get_audit_logger(
                        user=st.session_state.username,
                        role=st.session_state.get('role', 'N/A'),
                        action="SITEMAP_SCAN_ERROR",
                        details=f"Sitemap scan failed for {site_to_scan_url}: {str(e)}"
                    )
                finally:
                    st.session_state.sitemap_scan_in_progress = False
                    st.rerun() # Rerun to update UI based on scan results / error
            else:
                st.warning("Please enter a URL to scan.")

        # Displaying scan progress or results (Task 2.3 will enhance this)
        if st.session_state.sitemap_scan_in_progress:
            st.info("Sitemap scan in progress...") # This might be brief due to spinner + rerun
        elif st.session_state.sitemap_scan_completed:
            if st.session_state.sitemap_scan_error:
                st.error(f"Scan failed: {st.session_state.sitemap_scan_error}")
            elif st.session_state.discovered_sitemap_urls:
                st.success(f"Scan complete! Found {len(st.session_state.discovered_sitemap_urls)} URLs.")
                
                st.subheader("Select URLs for Scraping:")

                # Select All / Deselect All buttons
                col1, col2, col_spacer = st.columns([1,1,5])
                with col1:
                    if st.button("Select All", key="select_all_sitemap_urls"):
                        st.session_state.selected_sitemap_urls = set(st.session_state.discovered_sitemap_urls)
                        st.rerun()
                with col2:
                    if st.button("Deselect All", key="deselect_all_sitemap_urls"):
                        st.session_state.selected_sitemap_urls = set()
                        st.rerun()

                # Use a container for scrollability if many URLs
                # Max height for the checkbox list container
                max_height = "300px" 
                # Create a string for inline CSS style for the container
                container_style = f"max-height: {max_height}; overflow-y: auto; border: 1px solid #ccc; padding: 10px; border-radius: 5px;"
                
                # Markdown to inject CSS for the container class
                # st.markdown(f'''
                # <style>
                # .checkbox-container {{
                #     max-height: {max_height};
                #     overflow-y: auto;
                #     border: 1px solid #ccc;
                #     padding: 10px;
                #     border-radius: 5px;
                # }}
                # </style>
                # ''', unsafe_allow_html=True)
                # Instead of injecting CSS, use st.container() and manage checkboxes within it.
                # However, st.container() itself doesn't directly support height and overflow styling like a div.
                # A common workaround is to display a limited number or use st.expander for very long lists.
                # For now, let's list them. If too many, consider pagination or virtualization later.

                # For simplicity, we will list all. Streamlit will handle scroll if the main page area becomes too long.
                # Or, we can show a limited number with a message.
                # Let's show all checkboxes. Streamlit might make the section scrollable if content overflows page.
                # We will manage selections directly by updating the set on checkbox interaction.

                for i, url in enumerate(st.session_state.discovered_sitemap_urls):
                    # Define a callback function for this specific checkbox
                    def checkbox_changed(url_to_toggle):
                        if url_to_toggle in st.session_state.selected_sitemap_urls:
                            st.session_state.selected_sitemap_urls.remove(url_to_toggle)
                        else:
                            st.session_state.selected_sitemap_urls.add(url_to_toggle)
                    
                    # Check if the URL is currently selected
                    is_selected = url in st.session_state.selected_sitemap_urls
                    
                    # Create the checkbox
                    # Using a unique key by index or URL itself to ensure state is maintained across reruns
                    st.checkbox(
                        url, 
                        value=is_selected, 
                        key=f"sitemap_url_cb_{i}_{url}", # Unique key
                        on_change=checkbox_changed, 
                        args=(url,)
                    )
                
                if st.session_state.discovered_sitemap_urls:
                     st.caption(f"{len(st.session_state.selected_sitemap_urls)} / {len(st.session_state.discovered_sitemap_urls)} URLs selected.")

            else: # Scan completed, no error, but no URLs
                st.info("Scan completed. No URLs were found from sitemaps for the provided site.")

        st.markdown("**Option 4b: Crawl and Scrape Starting from URL**")
        crawl_start_url = st.text_input(
            "Starting URL for Crawl:",
            key="crawl_start_url_input",
            placeholder="https://example.com/startpage"
        )
        crawl_limit = st.number_input(
            "Max Pages to Scrape (Crawl Limit):",
            min_value=1,
            max_value=50, # Set a reasonable upper limit
            value=5, # Default limit
            step=1,
            key="crawl_limit_input",
            help="Maximum number of pages to scrape during the crawl, starting from the URL above."
        )
        st.markdown("---")
        # --- End Crawl Input Section ---

        # Report Generation
        st.subheader("5. Generate Report")
        if st.button("Generate Unified Report", key="generate_unified_report_button"):
            # Check if any input is provided
            if not (
                research_query or 
                st.session_state.processed_documents_content or 
                submitted_urls or 
                crawl_start_url or 
                st.session_state.get("selected_sitemap_urls")
            ):
                st.warning("Please provide a research query, upload documents, enter specific URLs, provide a starting URL to crawl, or select URLs from a site scan.")
            else:
                # Clear previous report content and results
                st.session_state.unified_report_content = ""
                st.session_state.scraped_web_content = [] 
                st.session_state.crawled_web_content = []

                # Determine which URLs to use for specific scraping and set up logging info
                urls_to_use_for_specific_scraping = []
                scrape_source_log_message = "No specific URLs defined for scraping."
                log_links_for_scraping = []

                if st.session_state.get("selected_sitemap_urls"): # Priority 1
                    urls_to_use_for_specific_scraping = list(st.session_state.selected_sitemap_urls)
                    scrape_source_log_message = f"{len(urls_to_use_for_specific_scraping)} URLs selected from sitemap scan."
                    log_links_for_scraping = urls_to_use_for_specific_scraping[:]
                elif submitted_urls: # Priority 2 (from text area)
                    urls_to_use_for_specific_scraping = submitted_urls[:]
                    scrape_source_log_message = f"{len(urls_to_use_for_specific_scraping)} URLs provided via text input."
                    log_links_for_scraping = urls_to_use_for_specific_scraping[:]

                # Prepare details for the initial audit log
                processed_doc_names = [doc['name'] for doc in st.session_state.processed_documents_content]
                log_details_parts = [
                    f"Research Query: '{research_query}'" if research_query else "No research query.",
                    f"Files: {len(processed_doc_names)} ({(', '.join(processed_doc_names[:3]) + '...') if len(processed_doc_names) > 3 else ', '.join(processed_doc_names)})."
                ]
                
                current_action_links_for_log = list(log_links_for_scraping) # Start with specific URLs if any

                if urls_to_use_for_specific_scraping:
                    log_details_parts.append(f"Specific URLs Source: {scrape_source_log_message}")
                    if crawl_start_url: # If specific URLs are used, note that crawl (if input) is skipped
                        log_details_parts.append(f"Crawl from '{crawl_start_url}': Skipped due to specific URLs being processed.")
                elif crawl_start_url: # No specific URLs, but crawl URL is provided
                    log_details_parts.append(f"Crawl: From '{crawl_start_url}', Limit: {crawl_limit}.")
                    current_action_links_for_log.append(f"[CRAWL_START] {crawl_start_url}")
                else: # No specific URLs and no crawl URL
                    log_details_parts.append("Web Content: No specific URLs or crawl initiated.")

                final_details_str_for_log = " ".join(log_details_parts)

                get_audit_logger(
                    user=st.session_state.username,
                    role=st.session_state.get('role','N/A'), 
                    action="REPORT_GENERATION_INITIATED",
                    details=final_details_str_for_log,
                    links=current_action_links_for_log if current_action_links_for_log else None,
                    model=st.session_state.get("selected_model", "N/A")
                )

                with st.spinner("Processing inputs and generating report..."):
                    # --- Stage 1a: Scrape Specific URLs (from sitemap selection or text area) ---
                    if urls_to_use_for_specific_scraping:
                        # Use the more descriptive scrape_source_log_message for UI feedback
                        ui_scrape_message_short = scrape_source_log_message.split('.')[0] # e.g., "X URLs selected from sitemap scan"
                        st.info(f"Starting web scraping for {len(urls_to_use_for_specific_scraping)} specific URL(s) ({ui_scrape_message_short})...")
                        scraped_data_specific = await process_urls(urls_to_use_for_specific_scraping, st.session_state.firecrawl_client)
                        st.session_state.scraped_web_content = scraped_data_specific
                    else:
                        st.info("No specific URLs (from sitemap scan or text input) to scrape.")
                        
                    # --- Stage 1b: Crawl & Scrape Site (only if no specific URLs were processed and crawl_start_url is provided) ---
                    if crawl_start_url and not urls_to_use_for_specific_scraping:
                        st.info(f"No specific URLs provided/selected. Proceeding with site crawl from {crawl_start_url}...")
                        crawled_data = await crawl_and_scrape_site(crawl_start_url, crawl_limit, st.session_state.firecrawl_client)
                        st.session_state.crawled_web_content = crawled_data
                    elif crawl_start_url and urls_to_use_for_specific_scraping:
                        st.info(f"Site crawl from '{crawl_start_url}' was skipped because specific URLs were selected/provided.")
                    elif not crawl_start_url and not urls_to_use_for_specific_scraping: # If no specific URLs and no crawl URL (but other inputs like docs exist)
                        st.info("No starting URL provided for site crawl, and no specific URLs selected/entered.")
                    
                    # --- Task 4: Combined Analysis & Report Generation --- 
                    st.info("Combining processed content and generating AI report...")
                    
                    # 1. Retrieve research query
                    # research_query variable is already available here

                    # 2. Retrieve text from processed documents
                    document_texts = []
                    if st.session_state.processed_documents_content:
                        for doc in st.session_state.processed_documents_content:
                            document_texts.append(f"--- Document: {doc['name']} ---\n{doc['text']}\n---")
                    combined_document_text = "\n".join(document_texts)

                    # 3. Retrieve text from BOTH scraped specific URLs and crawled URLs
                    scraped_content_texts = []
                    if st.session_state.scraped_web_content:
                        scraped_content_texts.append("--- Specific URLs Content ---")
                        for item in st.session_state.scraped_web_content:
                            if item["status"] == "success" and item.get("content"): 
                                scraped_content_texts.append(f"--- URL: {item['url']} ---\n{item['content']}\n---")
                    
                    crawled_content_texts = []
                    successful_crawls = [item for item in st.session_state.crawled_web_content if item["status"] == "success" and item.get("content")]
                    if successful_crawls:
                        crawled_content_texts.append("--- Crawled Site Content ---")
                        for item in successful_crawls:
                            crawled_content_texts.append(f"--- Crawled URL: {item['url']} ---\n{item['content']}\n---")
                                
                    combined_scraped_text = "\n".join(scraped_content_texts)
                    combined_crawled_text = "\n".join(crawled_content_texts)

                    # 4. Combine all text sources with the research query
                    # The research_query acts as the primary instruction/question.
                    # The system_prompt (in st.session_state.system_prompt) guides the AI's persona and output format.
                    
                    # Construct the main instruction based on whether a user query was provided
                    if research_query:
                        full_prompt_for_ai = f"Research Query: {research_query}\\n\\n"
                    else:
                        # Default instruction if query is empty - rely on system prompt and content
                        full_prompt_for_ai = "Research Goal: Please generate a comprehensive report based on the provided content (if any) and the overall objectives defined in the system prompt.\\n\\n"

                    # Append document content if available
                    if combined_document_text:
                        full_prompt_for_ai += f"Provided Document(s) Content:\\n{combined_document_text}\\n\\n"
                    else:
                        full_prompt_for_ai += "No documents were provided or processed.\\n\\n"
                        
                    # Append scraped web content if available
                    if combined_scraped_text:
                        full_prompt_for_ai += f"Provided Specific Web Content:\n{combined_scraped_text}\\n\\n"
                    # else: (No longer needed to say 'No web content' here, check combined below)
                    #    full_prompt_for_ai += "No specific web content was provided or successfully scraped.\\n\\n"
                        
                    if combined_crawled_text:
                        full_prompt_for_ai += f"Provided Crawled Web Content:\n{combined_crawled_text}\\n\\n"
                        
                    if not combined_scraped_text and not combined_crawled_text:
                         full_prompt_for_ai += "No web content was provided or successfully scraped/crawled.\\n\\n"
                    
                    full_prompt_for_ai += "Based on the research goal/query and all the provided content above, please generate a comprehensive report."

                    # 5. Call OpenRouterClient
                    # Ensure openrouter_client is initialized (it is, outside this button block)
                    
                    # --- DEBUGGING --- 
                    if st.session_state.get("role") == "admin": # Condition to show debug info only for admin
                        st.write("--- DEBUG INFO BEFORE AI CALL (Admin Only) ---")
                        # Prepare query snippet for debugging, handling potential quotes
                        debug_query_snippet = research_query[:50]
                        if len(research_query) > 50:
                            debug_query_snippet += "..."
                        st.write(f"Research Query Provided: {'Yes' if research_query else 'No'} (Content: '{debug_query_snippet}')") 
                        st.write(f"Combined Document Text Provided: {'Yes' if combined_document_text else 'No'} (Length: {len(combined_document_text)})")
                        st.write(f"Combined Scraped Text Provided: {'Yes' if combined_scraped_text else 'No'} (Length: {len(combined_scraped_text)})")
                        st.write(f"Combined Crawled Text Provided: {'Yes' if combined_crawled_text else 'No'} (Length: {len(combined_crawled_text)})")
                        st.write("--- END DEBUG INFO ---")
                    # --- END DEBUGGING ---

                    try:
                        # --- Get the selected model ---
                        model_to_use = st.session_state.get("selected_model", OPENROUTER_PRIMARY_MODEL) # Fallback just in case
                        # Remove the log here as it's logged earlier and below
                        # log_audit_event(st.session_state.username, st.session_state.get('role','N/A'), "MODEL_SELECTION_USED", f"Using model: {model_to_use} for report generation.")
                        # --- End Get model ---

                        # --- DEBUG PRINT FOR MAIN REPORT ---
                        print(f"DEBUG (Main Report): About to call LLM. Model identifier type: {type(model_to_use)}, value: {model_to_use}")
                        # --- END DEBUG PRINT FOR MAIN REPORT ---

                        ai_generated_report = await st.session_state.openrouter_client.generate_response(
                            prompt=full_prompt_for_ai, 
                            system_prompt=st.session_state.system_prompt,
                            model_override=model_to_use # Pass selected model for main report
                        )
                        if ai_generated_report:
                            st.session_state.unified_report_content = ai_generated_report
                            st.success("AI Report generated successfully!")
                            
                            st.session_state.report_generated_for_chat = True
                            report_id = f"report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S%f')}"
                            st.session_state.current_report_id_for_chat = report_id
                            st.session_state.current_chat_session_id = None 
                            st.session_state.chat_ui_expanded = False # Ensure chat is initially collapsed

                            # --- Build RAG Index for the new report --- 
                            with st.spinner(f"Preparing RAG context for report {report_id}..."):
                                embedding_model = get_embedding_model() # Uses cached model
                                
                                all_text_for_rag = []
                                if st.session_state.unified_report_content:
                                    all_text_for_rag.append(st.session_state.unified_report_content)
                                
                                for doc in st.session_state.get("processed_documents_content", []):
                                    all_text_for_rag.append(f"--- Document: {doc.get('name', 'Unnamed Document')} ---\n{doc.get('text', '')}")
                                
                                for item in st.session_state.get("scraped_web_content", []):
                                    if item.get("status") == "success" and item.get("content"):
                                        all_text_for_rag.append(f"--- Web Scrape: {item.get('url', 'Unknown URL')} ---\n{item.get('content', '')}")
                                
                                for item in st.session_state.get("crawled_web_content", []):
                                    if item.get("status") == "success" and item.get("content"):
                                        all_text_for_rag.append(f"--- Web Crawl: {item.get('url', 'Unknown URL')} ---\n{item.get('content', '')}")
                                
                                combined_text = "\n\n---\n\n".join(all_text_for_rag)
                                text_chunks = split_text_into_chunks(combined_text) # Uses defaults from rag_utils
                                
                                if text_chunks:
                                    faiss_index = build_faiss_index(text_chunks, embedding_model)
                                    if faiss_index:
                                        st.session_state.rag_contexts[report_id] = {
                                            "index": faiss_index,
                                            "chunks": text_chunks,
                                            "embedding_model_name": DEFAULT_EMBEDDING_MODEL # Use the imported constant
                                        }
                                        st.success(f"RAG context built for report {report_id} with {len(text_chunks)} chunks.")
                                    else:
                                        st.warning(f"Failed to build RAG index for report {report_id}. Chat may have limited context.")
                                        st.session_state.rag_contexts[report_id] = None # Indicate failure
                                else:
                                    st.warning(f"No text content found to build RAG index for report {report_id}.")
                                    st.session_state.rag_contexts[report_id] = None # Indicate failure
                            # --- End RAG Index Building ---

                            successfully_scraped_urls = [item['url'] for item in st.session_state.scraped_web_content if item["status"] == "success"]
                            query_for_log = research_query if research_query else '[SYSTEM PROMPT]'
                            success_details = f"AI report generated for query: '{query_for_log}'. Docs: {len(processed_doc_names)}. Scraped URLs: {len(successfully_scraped_urls)}."
                            
                            get_audit_logger( # Replaced log_audit_event
                                user=st.session_state.username,
                                role=st.session_state.get('role','N/A'), 
                                action="REPORT_GENERATION_SUCCESS", 
                                details=success_details,
                                links=successfully_scraped_urls if successfully_scraped_urls else None,
                                model=model_to_use # Log model used
                            )
                        else:
                            st.session_state.unified_report_content = "Failed to generate AI report. The AI returned an empty response."
                            st.error("AI report generation failed or returned empty.")
                            
                            # --- Additions for Chat UI (on failure) ---
                            st.session_state.report_generated_for_chat = False
                            st.session_state.current_report_id_for_chat = None
                            st.session_state.chat_ui_expanded = False # Ensure chat is not shown/expanded
                            # --- End Additions for Chat UI ---

                            query_for_log = research_query if research_query else '[SYSTEM PROMPT]'
                            get_audit_logger( # Replaced log_audit_event
                                user=st.session_state.username,
                                role=st.session_state.get('role','N/A'),
                                action="REPORT_GENERATION_FAILURE",
                                details=f"AI returned empty report for query: '{query_for_log}'"
                            )
                            get_audit_logger( # Replaced log_audit_event
                                user=st.session_state.username,
                                role=st.session_state.get('role','N/A'),
                                action="REPORT_GENERATION_FAILURE",
                                details=f"AI returned empty report for query: '{query_for_log}'",
                                model=model_to_use
                            )
                    except Exception as e:
                        st.session_state.unified_report_content = f"An error occurred during AI report generation: {str(e)}"
                        st.error(f"Error calling AI: {e}")

                        # --- Additions for Chat UI (on exception) ---
                        st.session_state.report_generated_for_chat = False
                        st.session_state.current_report_id_for_chat = None
                        st.session_state.chat_ui_expanded = False # Ensure chat is not shown/expanded
                        # --- End Additions for Chat UI ---

                        query_for_log = research_query if research_query else '[SYSTEM PROMPT]'
                        get_audit_logger( # Replaced log_audit_event
                            user=st.session_state.username,
                            role=st.session_state.get('role','N/A'),
                            action="REPORT_GENERATION_ERROR",
                            details=f"Error during AI call for query '{query_for_log}': {str(e)}"
                        )
                        get_audit_logger( # Replaced log_audit_event
                            user=st.session_state.username,
                            role=st.session_state.get('role','N/A'),
                            action="REPORT_GENERATION_ERROR",
                            details=f"Error during AI call for query '{query_for_log}': {str(e)}",
                            model=model_to_use
                        )
                    
                    # Remove the debug JSON output now that real processing is in place
                    # st.markdown("--- DEBUG: Information Collected ---\")
                    # st.json({
                    #     "research_query": research_query,
                    #     "processed_documents": st.session_state.processed_documents_content,
                    #     "scraped_urls_content": st.session_state.scraped_web_content
                    # })
                    # st.info("Report generation logic (Task 4) is pending. Displaying collected data for now.")
                    # --- End of Task 4 ---
                
                # Rerun to display the generated report or messages
                st.rerun()
        
        # Area for actual report display (will be populated by backend logic)
        if "unified_report_content" in st.session_state and st.session_state.unified_report_content:
            st.markdown("---")
            st.subheader("Generated Report")
            st.markdown(st.session_state.unified_report_content)
            st.download_button(
                label="Download Unified Report",
                data=st.session_state.unified_report_content,
                file_name="unified_research_report.md",
                mime="text/markdown",
                key="download_actual_unified_report",
                on_click=lambda: get_audit_logger( # Replaced log_audit_event
                    user=st.session_state.username, 
                    role=st.session_state.get('role', 'N/A'), 
                    action="REPORT_DOWNLOADED", 
                    details=f"User downloaded report: unified_research_report.md for query '{(st.session_state.get('research_query_input', 'QUERY_NOT_IN_SESSION_FOR_DOWNLOAD_LOG'))}'"
                )
            )
        # elif "generate_unified_report_button_clicked" in st.session_state and \
        #      st.session_state.generate_unified_report_button_clicked and \
        #      not st.session_state.unified_report_content:
        #      # This logic might be implicitly handled by the spinner and rerun, 
        #      # or if a specific message for "no report generated" is desired, it can be placed here.
        #      # For now, if content is empty, nothing is shown, which is acceptable.
        #      pass 

        st.markdown("---") # Separator before Admin Panel
        
        # ==== ADMIN PANEL ====
        if st.session_state.get("role") == "admin":
            st.header("Admin Panel - Audit Logs")
            
            log_file_path = AUDIT_LOG_FILE_PATH 
            st.write(f"DEBUG: Checking for log file at: {log_file_path.resolve()}") # DEBUG
            log_data = []

            if log_file_path.exists():
                st.write(f"DEBUG: Log file exists. File size: {log_file_path.stat().st_size} bytes") # DEBUG
                try:
                    with open(log_file_path, 'r') as f:
                        # Read lines in reverse to show newest first
                        log_lines = f.readlines() # Read all lines first
                        st.write(f"DEBUG: Read {len(log_lines)} lines from log file.") # DEBUG
                        for line in reversed(log_lines): # Then reverse
                            parsed = parse_log_line(line)
                            if parsed:
                                log_data.append(parsed)
                                
                    if log_data:
                        df = pd.DataFrame(log_data)
                        # Reorder columns for better readability
                        cols_order = ["Timestamp", "Username", "Role", "Action", "Model", "Details", "Links"]
                        cols_to_display = [col for col in cols_order if col in df.columns]
                        
                        # Configure column widths
                        column_config = {
                            "Links": st.column_config.TextColumn("Links", width="medium"),
                            "Details": st.column_config.TextColumn("Details", width="medium"),
                             # Set timestamp width if needed
                            "Timestamp": st.column_config.TextColumn("Timestamp", width="small"),
                        }
                        # Apply config only for columns present in the dataframe
                        active_column_config = {k: v for k, v in column_config.items() if k in cols_to_display}
                        
                        st.dataframe(df[cols_to_display], column_config=active_column_config)
                    else:
                        st.info("Audit log file exists but contains no parseable entries (or all entries were unparseable).") # Modified message
                        
                except Exception as e:
                    st.error(f"Error reading or parsing audit log file: {e}")
                    get_audit_logger( # Replaced log_audit_event
                        user=st.session_state.username,
                        role=st.session_state.role,
                        action="AUDIT_LOG_READ_ERROR",
                        details=f"Admin panel failed to read/parse audit log: {e}"
                    )
            else:
                st.warning("Audit log file not found. Logging may not be configured or no events logged yet.")
                st.write(f"DEBUG: Attempted to access {log_file_path.resolve()} but it was not found.") # DEBUG
            
            if st.button("Refresh Logs", key="refresh_logs_button"):
                 st.rerun() # Simple way to refresh the view by rerunning the script
        # ==== END ADMIN PANEL ====

        # --- Display Chat Interface ---
        # This should be called conditionally after a report is generated.
        # For example, if a report is displayed using st.markdown(unified_report_markdown),
        # then call display_chat_interface() right after it.

        if st.session_state.get("report_generated_for_chat", False) and st.session_state.get("current_report_id_for_chat"):
            await display_chat_interface_v2()
        
    else: # Not authenticated
        st.info("Please log in or create an account to access the research pipeline.")
        # The login/signup logic from the sidebar already handles this visibility.
        # This message is a fallback or main page content for unauthenticated users.

async def run_main():
    await main()

if __name__ == "__main__":
    asyncio.run(run_main()) 